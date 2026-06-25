import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = 'Times New Roman'
    heading.style.font.color.rgb = None # Black

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(10)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(9)
    return p

def main():
    SAVE_DIR = "saved_models"
    doc = Document()

    # --- TITLE ---
    title = doc.add_heading('Federated Healthcare AI Framework for Diabetes Prediction using BiLSTM and CNN–BiLSTM–Attention', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "Member 4 Contribution Report\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. ABSTRACT ---
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, 
        "Healthcare institutions generate vast amounts of sensitive patient data, making collaborative deep learning difficult due to strict privacy regulations such as HIPAA. This paper presents a privacy-preserving Federated Healthcare AI Framework for Diabetes Prediction. Focusing specifically on sequential and hybrid deep learning architectures, we implement and evaluate a Bidirectional Long Short-Term Memory (BiLSTM) network and a hybrid CNN–BiLSTM–Attention model. The models are trained on a clinical diabetes dataset utilizing essential markers such as HbA1c and blood glucose levels. To ensure privacy, the best-performing architecture was deployed across four simulated hospital clients using the Federated Averaging (FedAvg) algorithm."
    )

    # --- 2. INTRODUCTION ---
    add_heading(doc, '1. Introduction', level=1)
    add_paragraph(doc, 
        "Early and accurate prediction of diabetes is vital for effective clinical intervention. While Deep Learning (DL) has shown immense promise in healthcare analytics, the centralization of Electronic Health Records (EHR) poses severe privacy risks.\n\n"
        "This implementation utilizes Federated Learning (FL), allowing models to train on decentralized data without transferring raw patient information. We explore two specific DL architectures for this tabular clinical data: a traditional BiLSTM to capture feature-wise dependencies, and a Hybrid CNN–BiLSTM–Attention model designed to extract spatial hierarchies, process dependencies, and dynamically weight the most critical clinical features."
    )

    # --- 3. METHODOLOGY ---
    add_heading(doc, '2. Methodology', level=1)
    
    add_heading(doc, '2.1 Data Preprocessing', level=2)
    add_paragraph(doc, 
        "The dataset was dynamically preprocessed to ensure mathematical compatibility with neural network tensors. Categorical variables were transformed using one-hot encoding to prevent runtime conversion errors. Features were normalized using Scikit-Learn’s StandardScaler fitted only on the training set to prevent data leakage. Finally, balanced class weights were calculated to penalize majority-class misclassifications."
    )

    add_heading(doc, '2.2 Traditional Model: BiLSTM', level=2)
    add_paragraph(doc, 
        "The BiLSTM architecture processes tabular patient data by reshaping the normalized feature vector into a timestep sequence. The network operates bidirectionally, evaluating clinical features both forwards and backwards to capture holistic dependencies. The architecture consists of:\n"
        "• Input Reshape Layer (1, Features)\n"
        "• Bidirectional LSTM Layer (64 units, return_sequences=False)\n"
        "• Dropout Layer (Rate: 0.3) for regularization\n"
        "• Dense Classifier (32 units, ReLU activation)\n"
        "• Dense Output Layer (2 units, Softmax activation)"
    )

    add_heading(doc, '2.3 Hybrid Model: CNN + BiLSTM + Attention', level=2)
    add_paragraph(doc, 
        "The hybrid model leverages multiple deep learning paradigms to maximize feature extraction from the clinical data. A Conv1D layer (64 filters, kernel size 3) followed by MaxPooling1D extracts local clinical feature groupings. A Bidirectional LSTM interprets the pooled feature maps. A native Keras Attention layer calculates alignment scores between LSTM hidden states, dynamically assigning higher mathematical weights to dominant clinical features. Finally, GlobalAveragePooling1D flattens the context vector for the Dense classifier."
    )

    # --- 4. ARCHITECTURE & WORKFLOW DIAGRAMS ---
    add_heading(doc, '3. Architecture and Workflow', level=1)
    add_paragraph(doc, "Figure 1: Hybrid CNN-BiLSTM-Attention Architecture Pipeline")
    add_code_block(doc, 
        "[Input: Clinical Patient Features]\n"
        "          ↓\n"
        "[Reshape: (Features, 1)]\n"
        "          ↓\n"
        "[Conv1D: 64 filters, kernel=3, ReLU]\n"
        "          ↓\n"
        "[MaxPooling1D: pool_size=2]\n"
        "          ↓\n"
        "[Bidirectional LSTM: 64 units, return_seq=True]\n"
        "          ↓\n"
        "[Attention Layer: Self-Alignment Scoring]\n"
        "          ↓\n"
        "[GlobalAveragePooling1D]\n"
        "          ↓\n"
        "[Dense: 32 units, ReLU] + [Dropout: 0.3]\n"
        "          ↓\n"
        "[Dense: 2 units, Softmax]\n"
        "          ↓\n"
        "[Output: Diabetic / Non-Diabetic Probability]"
    )

    # --- 5. IMPLEMENTATION DETAILS ---
    add_heading(doc, '4. Implementation Details', level=1)
    add_paragraph(doc, 
        "The project was implemented strictly using Python 3 and TensorFlow/Keras. The centralized models utilized the Adam Optimizer and Sparse Categorical Crossentropy loss function, training for 15 epochs with a batch size of 32.\n\n"
        "Hardware Optimization: To bypass recurrent network deadlocks on Apple Silicon GPUs (Metal), explicit environment variables (TF_CPP_MIN_LOG_LEVEL) were utilized alongside tf.config.set_visible_devices([], 'GPU') to force stable CPU execution for complex recurrent matrix operations."
    )

    # --- 6. FEDERATED LEARNING FRAMEWORK ---
    add_heading(doc, '5. Federated Learning Framework (FedAvg)', level=1)
    add_paragraph(doc, 
        "A Federated Averaging (FedAvg) pipeline was engineered to simulate a decentralized network of four distinct hospital nodes (Hospitals A, B, C, D). "
    )
    add_heading(doc, '5.1 Federated Workflow', level=2)
    add_paragraph(doc, 
        "1. Initialization: The central server initializes the global model using the best-performing architecture from the centralized baseline phase.\n"
        "2. Distribution: The global weights are transmitted to the participating hospital nodes.\n"
        "3. Local Training: Each hospital trains the model locally on its isolated Electronic Health Records (EHR) for 1 local epoch. Raw data is strictly maintained on-premises.\n"
        "4. Aggregation: The central server collects the updated weights from the clients and aggregates them using an arithmetic mean to form the new global model.\n"
        "5. Iteration: This process was repeated across 3 communication rounds, utilizing fractional client selection (2 random hospitals per round) to optimize computational overhead."
    )

    # --- 7. COMPARATIVE ANALYSIS ---
    add_heading(doc, '6. Comparative Analysis', level=1)
    
    metrics_file = os.path.join(SAVE_DIR, "model_comparison.csv")
    if os.path.exists(metrics_file):
        df = pd.read_csv(metrics_file)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Accuracy'
        hdr_cells[2].text = 'Precision'
        hdr_cells[3].text = 'Recall'
        hdr_cells[4].text = 'F1-Score'

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Model'])
            row_cells[1].text = f"{row.get('Accuracy', 0):.4f}"
            row_cells[2].text = f"{row.get('Precision', 0):.4f}"
            row_cells[3].text = f"{row.get('Recall', 0):.4f}"
            row_cells[4].text = f"{row.get('F1 Score', 0):.4f}"
            
        add_paragraph(doc, "\nTable 1: Model Performance Comparison based on execution metrics.")
    else:
        add_paragraph(doc, "[model_comparison.csv not found in saved_models directory. Table omitted.]")

    # --- 8. VISUALIZATIONS (Optional Images) ---
    add_heading(doc, '7. Visualizations', level=1)
    add_paragraph(doc, "The implemented framework automatically generated accuracy progression tracking and confusion matrices to validate model stability.")
    
    # Try to embed images if they exist locally
    image_list = [
        ("member4_cnn_bilstm_attention_roc_pr.png", "Figure 2: ROC and Precision-Recall Curves for Hybrid Architecture."),
        ("federated_progression.png", "Figure 3: Global Accuracy Progression across FedAvg Rounds.")
    ]

    images_added = False
    for img_file, caption in image_list:
        img_path = os.path.join(SAVE_DIR, img_file)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.0))
            add_paragraph(doc, caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            images_added = True
            
    if not images_added:
        add_paragraph(doc, "[Images will be automatically embedded here if present in the saved_models directory.]")

    # --- 9. CONCLUSION ---
    add_heading(doc, '8. Conclusion', level=1)
    add_paragraph(doc, 
        "This project successfully implemented a privacy-preserving healthcare AI framework for diabetes prediction. The Hybrid CNN–BiLSTM–Attention architecture demonstrated robust feature extraction capabilities on tabular clinical data. Furthermore, the successful deployment of the FedAvg algorithm proved that highly accurate, decentralized deep learning models can be trained collaboratively across medical institutions without transmitting sensitive patient records."
    )

    # --- SAVE ---
    report_path = "Member4_Full_IEEE_Report.docx"
    doc.save(report_path)
    print(f"✅ Success! Document generated and saved to: {report_path}")

if __name__ == "__main__":
    main()